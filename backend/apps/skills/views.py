import re
import io
import json
import os
from docx import Document
from pypdf import PdfReader, PdfWriter
from PIL import Image, ImageDraw, ImageFont
from django.core.files.base import ContentFile
from groq import Groq
from rest_framework import viewsets, status, generics, permissions
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.views import APIView

from .models import Skill, SkillProficiency, UserSkill, RequiredSkill, SkillGapReport, ResumeData, ResumeHistory, ResumeBuilderProfile, CustomTemplate, GeneratedResume
from .serializers import SkillSerializer, UserSkillSerializer, SkillGapReportSerializer, ResumeDataSerializer, ResumeHistorySerializer, ResumeBuilderProfileSerializer, CustomTemplateSerializer, GeneratedResumeSerializer
from .analysis import analyze_skills, extract_skills_from_text, calculate_match_score, generate_roadmap, extract_text_from_file

class SkillViewSet(viewsets.ViewSet):
    permission_classes = [AllowAny]

    def list(self, request):
        skills = Skill.objects.all()
        serializer = SkillSerializer(skills, many=True)
        return Response(serializer.data)

    def retrieve(self, request, pk=None):
        skill = Skill.objects.get(pk=pk)
        serializer = SkillSerializer(skill)
        return Response(serializer.data)

    def analyze(self, request):
        skill_data = request.data.get('skills', [])
        analysis_result = analyze_skills(skill_data)
        return Response(analysis_result)


class SkillListCreateView(generics.ListCreateAPIView):
    serializer_class = SkillSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        profs = SkillProficiency.objects.filter(user=self.request.user)
        skill_ids = profs.values_list('skill_id', flat=True)
        return Skill.objects.filter(id__in=skill_ids)

    def perform_create(self, serializer):
        name = serializer.validated_data.get('name')
        description = serializer.validated_data.get('description', '')
        skill, _ = Skill.objects.get_or_create(name=name, defaults={'description': description})
        level = self.request.data.get('level', '')
        years = self.request.data.get('years_of_experience', 0)
        UserSkill.objects.create(skill=skill, user=self.request.user, level=level or 'beginner', category=self.request.data.get('category','other'), progress_percentage=self.request.data.get('progress_percentage',0))
        return skill

class UserSkillViewSet(viewsets.ModelViewSet):
    serializer_class = UserSkillSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return self.request.user.skills_owned.select_related('skill')

    def perform_create(self, serializer):
        serializer.save()

    @action(detail=False, methods=['get'])
    def me(self, request):
        qs = self.get_queryset()
        serializer = self.get_serializer(qs, many=True)
        return Response(serializer.data)


class ResumeUploadView(generics.CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ResumeDataSerializer

    def post(self, request):
        file_obj = request.FILES.get('resume')
        if not file_obj:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)
        
        text = extract_text_from_file(file_obj)
        if not text:
            return Response({"error": "Failed to extract text from file"}, status=status.HTTP_400_BAD_REQUEST)
            
        skills = extract_skills_from_text(text)
        
        resume_data, created = ResumeData.objects.update_or_create(
            user=request.user,
            defaults={
                'extracted_skills': skills,
                'resume_file': file_obj
            }
        )
        
        # Save to history
        ResumeHistory.objects.create(
            user=request.user,
            file_name=file_obj.name,
            resume_file=file_obj,
            ats_score=min(100, 60 + (len(skills) * 2)),
            job_match_score=75,
            extracted_data={"skills": skills}
        )
        
        skill_count = len(skills)
        strength = [
            {'label': 'Technical Skills', 'value': min(95, skill_count * 8)},
            {'label': 'Keyword Match', 'value': 75},
            {'label': 'Format & Structure', 'value': 90}
        ]
        
        suggestions = [
            "Add more specific metrics to your experience bullet points.",
            f"You have {skill_count} key skills identified.",
            "Standardize your date formats across all experience entries."
        ]
        
        return Response({
            "message": "Resume processed successfully",
            "skills": skills,
            "id": resume_data.id,
            "strength_breakdown": strength,
            "ai_suggestions": suggestions,
            "readiness_score": min(100, 60 + (skill_count * 2))
        })

class CheckResumeStatusView(generics.RetrieveAPIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            resume = request.user.resume_data
            skills = resume.extracted_skills or []
            skill_count = len(skills)
            
            strength = [
                {'label': 'Technical Skills', 'value': min(95, skill_count * 8)},
                {'label': 'Keyword Match', 'value': 75},
                {'label': 'Format & Structure', 'value': 90}
            ]
            
            suggestions = [
                "Add more specific metrics to your experience bullet points.",
                f"You have {skill_count} key skills identified.",
                "Ensure your contact details and LinkedIn profile are up to date."
            ]

            return Response({
                "exists": True,
                "skills": skills,
                "created_at": resume.created_at,
                "strength_breakdown": strength,
                "readiness_score": min(100, 60 + (skill_count * 2)),
                "ai_suggestions": suggestions
            })
        except ResumeData.DoesNotExist:
            return Response({"exists": False})

class AnalyzeSkillGapView(generics.CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = SkillGapReportSerializer

    def post(self, request):
        role = request.data.get('role', 'Software Developer')
        jd_text = request.data.get('job_description', '')
        exp_level = request.data.get('experience_level', 'Mid')
        industry = request.data.get('industry', 'Tech')
        manual_skills = request.data.get('manual_skills', [])
        
        user_skills = []
        try:
            resume = request.user.resume_data
            user_skills = list(resume.extracted_skills)
        except ResumeData.DoesNotExist:
            pass

        user_skill_levels = {s.lower(): 3 for s in user_skills}
        for ms in manual_skills:
            lvl_map = {'Beginner': 2, 'Intermediate': 4, 'Advanced': 5}
            user_skill_levels[ms['name'].lower()] = lvl_map.get(ms['level'], 2)

        if jd_text:
            req_names = extract_skills_from_text(jd_text)
            required_skills_objs = RequiredSkill.objects.filter(skill_name__in=[s for s in req_names])
            req_list = []
            for name in req_names:
                existing = required_skills_objs.filter(skill_name__iexact=name).first()
                req_list.append({
                    'name': name,
                    'importance': existing.importance_level if existing else 3,
                    'category': existing.category if existing else 'technical'
                })
        else:
            required_skills_objs = RequiredSkill.objects.filter(role_name__icontains=role)
            if not required_skills_objs.exists():
                mock_required = [('React', 4, 'technical'), ('TypeScript', 3, 'technical'), ('Node.js', 3, 'technical'), ('Problem Solving', 3, 'soft')]
                RequiredSkill.objects.bulk_create([
                    RequiredSkill(role_name=role, skill_name=s, importance_level=l, category=c)
                    for s, l, c in mock_required
                ])
                required_skills_objs = RequiredSkill.objects.filter(role_name=role)
            
            req_list = [{
                'name': rs.skill_name,
                'importance': rs.importance_level,
                'category': rs.category
            } for rs in required_skills_objs]

        strong_match = []
        partial_match = []
        missing = []
        
        tech_score_val = 0
        tech_count = 0
        soft_score_val = 0
        soft_count = 0

        for req in req_list:
            req_name_l = req['name'].lower()
            user_lvl = user_skill_levels.get(req_name_l, 0)
            target_lvl = 3
            if exp_level == 'Mid': target_lvl = 4
            if exp_level == 'Senior': target_lvl = 5

            skill_info = {
                'name': req['name'],
                'required': req['importance'],
                'user_level': user_lvl,
                'category': req['category']
            }

            if user_lvl >= target_lvl:
                strong_match.append(skill_info)
            elif user_lvl > 0:
                partial_match.append(skill_info)
            else:
                missing.append(skill_info)

            weight = req['importance']
            match_ratio = min(1.0, user_lvl / target_lvl) if user_lvl > 0 else 0
            
            if req['category'] == 'technical':
                tech_score_val += (match_ratio * weight)
                tech_count += weight
            else:
                soft_score_val += (match_ratio * weight)
                soft_count += weight

        readiness = ((tech_score_val + soft_score_val) / (tech_count + soft_count) * 100) if (tech_count + soft_count) > 0 else 0
        t_score = (tech_score_val / tech_count * 100) if tech_count > 0 else 0
        s_score = (soft_score_val / soft_count * 100) if soft_count > 0 else 0

        priority_list = [m['name'] for m in missing[:5]]
        priority_sk = priority_list[0] if priority_list else ""
        impact_boost = 15 if missing else 5
        
        target_score = min(100, int(readiness) + impact_boost)
        if priority_list:
            ai_summary = f"Your profile is a solid match, but there are {len(missing)} key gaps. If you improve in {', '.join(priority_list[:2])}, your match score could reach {target_score}%."
        else:
            ai_summary = "You have a very strong coverage of required skills."

        report = SkillGapReport.objects.create(
            user=request.user,
            role_name=role,
            matched_skills=strong_match + partial_match,
            missing_skills=missing,
            readiness_score=round(readiness, 1),
            technical_score=round(t_score, 1),
            soft_skill_score=round(s_score, 1)
        )

        try:
            profile = request.user.profile
            profile.job_readiness_score = int(readiness)
            profile.dream_job = role
            profile.save()
            profile.calculate_completeness()
        except: pass

        response_data = SkillGapReportSerializer(report).data
        response_data.update({
            'strong_matches': strong_match,
            'partial_matches': partial_match,
            'missing_skills': missing,
            'ai_insight': ai_summary,
            'roadmap': generate_roadmap(missing),
            'priority_learning': priority_sk,
            'priority_skills': priority_list,
            'improvement_impact': impact_boost
        })
        
        return Response(response_data)


class RecommendLearningView(generics.CreateAPIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        role = request.data.get('role', 'Developer')
        missing = request.data.get('missing_skills', [])
        recommendations = {
            'courses': [
                {
                    'name': f'Mastering {s["name"]}', 
                    'platform': 'Coursera', 
                    'level': 'Intermediate',
                    'url': f'https://www.coursera.org/courses?query={s["name"].replace(" ", "%20")}'
                }
                for s in missing if s['category'] == 'technical'
            ][:3],
            'projects': [
                {'name': f'Building a Real-world {role} App', 'difficulty': 'Advanced'},
                {'name': f'Next-gen {role} Toolkit', 'difficulty': 'Intermediate'}
            ],
            'roadmap': generate_roadmap(missing),
            'estimated_time': f"{max(4, len(missing) * 2)} weeks"
        }
        return Response(recommendations)


class ExtractSkillsView(generics.CreateAPIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        text = request.data.get('text', '')
        if not text:
            return Response({'skills': []})
        skills = extract_skills_from_text(text)
        return Response({'skills': skills})

class ResumeHistoryViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = ResumeHistorySerializer
    def get_queryset(self): return ResumeHistory.objects.filter(user=self.request.user)

from rest_framework.decorators import action

class ResumeBuilderViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = ResumeBuilderProfileSerializer

    def get_queryset(self):
        return ResumeBuilderProfile.objects.filter(user=self.request.user)

    def get_object(self):
        obj, _ = ResumeBuilderProfile.objects.get_or_create(user=self.request.user)
        return obj

    @action(detail=False, methods=['get', 'put', 'patch'])
    def current(self, request):
        obj = self.get_object()
        if request.method == 'GET':
            serializer = self.get_serializer(obj)
            return Response(serializer.data)
        elif request.method in ['PUT', 'PATCH']:
            # For simplicity, we'll allow partial updates even for PUT in this context
            # We want to map the incoming JSON data to the model fields
            data = request.data
            obj.personal_details = data.get('personal_details', obj.personal_details)
            obj.summary = data.get('summary', obj.summary)
            obj.skills = data.get('skills', obj.skills)
            obj.experience = data.get('experience', obj.experience)
            obj.education = data.get('education', obj.education)
            obj.projects = data.get('projects', obj.projects)
            obj.certifications = data.get('certifications', obj.certifications)
            obj.achievements = data.get('achievements', obj.achievements)
            obj.template_config = data.get('template_config', obj.template_config)
            obj.save()
            serializer = self.get_serializer(obj)
            return Response(serializer.data)


class ResumeAIImproveView(APIView):
    permission_classes = [IsAuthenticated]
    def post(self, request):
        action = request.data.get('action', 'improve')
        content = request.data.get('content', '')
        role = request.data.get('role', 'Professional')
        api_key = os.environ.get('GROQ_API_KEY')
        client = Groq(api_key=api_key)
        prompts = {
            'improve': f'Rewrite this resume content for a {role} role to be more impact-oriented and professional: {content}',
            'summary': f'Generate a professional resume summary for a {role} with these skills: {content}',
            'keywords': f'Suggest 10 high-impact keywords for a {role} resume based on these skills: {content}'
        }
        prompt = prompts.get(action, prompts['improve'])
        try:
            completion = client.chat.completions.create(model='llama-3.1-8b-instant', messages=[{'role': 'user', 'content': prompt}])
            return Response({'result': completion.choices[0].message.content})
        except Exception as e: return Response({'error': str(e)}, status=500)

class ATSInsightView(APIView):
    """
    Resume Intelligence 2.0 Engine.
    Performs dynamic ATS scoring by matching resume content against 
    target job requirements (RequiredSkill models).
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user = request.user
        profile = getattr(user, 'profile', None)
        target_job = profile.dream_job if profile else "Software Engineer"
        
        try:
            resume = user.resume_data
            text = extract_text_from_file(resume.resume_file)[:4000].lower()
        except:
            return Response({'error': 'No resume found. Please upload a resume first.'}, status=400)

        # 1. Fetch Job Requirements
        job_skills = RequiredSkill.objects.filter(role_name__icontains=target_job)
        if not job_skills.exists():
            # Fallback to general Software Engineer if specific job not found
            job_skills = RequiredSkill.objects.filter(role_name__icontains="Software Engineer")

        # 2. Match Skills
        matches = []
        gaps = []
        heatmap = {
            'technical': {'match': 0, 'total': 0},
            'soft': {'match': 0, 'total': 0},
            'tools': {'match': 0, 'total': 0}
        }

        for skill in job_skills:
            category = skill.category if skill.category in heatmap else 'technical'
            heatmap[category]['total'] += 1
            
            # Simple word matching (can be improved with NLP/Embeddings)
            if skill.skill_name.lower() in text:
                matches.append(skill.skill_name)
                heatmap[category]['match'] += 1
            else:
                gaps.append({
                    'name': skill.skill_name,
                    'importance': skill.importance_level,
                    'impact': min(20, skill.importance_level * 4)
                })

        # 3. Calculate Scores
        keyword_score = int((len(matches) / job_skills.count() * 100)) if job_skills.count() > 0 else 70
        
        # Formatting Score (Basic heuristic)
        fmt_score = 90
        if len(text) < 1000: fmt_score -= 20 # Too short
        if 'experience' not in text: fmt_score -= 15
        if 'education' not in text: fmt_score -= 10
        if '@' not in text: fmt_score -= 15 # Missing contact?

        readability = 85 # Placeholder for legibility

        # 4. Generate Keyword Injections
        # Focus on Critical (4) or Mandatory (5) gaps
        injections = []
        critical_gaps = [g for g in gaps if g['importance'] >= 4]
        for g in critical_gaps[:3]:
            injections.append(f"Inject '{g['name']}' into your Experience section to boost your ATS match by ~{g['impact']}%")

        # 5. Build Final Response
        return Response({
            'ats_score': int((keyword_score * 0.7) + (fmt_score * 0.3)),
            'keyword_score': keyword_score,
            'formatting_score': fmt_score,
            'readability_score': readability,
            'target_job': target_job,
            'heatmap': {
                k: int((v['match']/v['total']*100)) if v['total'] > 0 else 0 
                for k, v in heatmap.items()
            },
            'keyword_injections': injections,
            'missing_critical': [g['name'] for g in critical_gaps],
            'suggestions': [
                'Quantify your achievements (e.g., Improved X by Y%)' if 'resposible for' in text else 'Well structured achievements',
                f'Add missing critical skills: {", ".join([g["name"] for g in critical_gaps[:2]])}' if critical_gaps else 'Core technical keywords are well represented',
                'Ensure your email and LinkedIn are clearly visible' if '@' not in text else 'Contact information detected'
            ]
        })

class ResumeExportView(APIView):
    permission_classes = [IsAuthenticated]
    def post(self, request):
        # Mock export
        return Response({'message': 'Export successful', 'url': '/mock-download.pdf'})

class CommunityTemplateViewSet(viewsets.ModelViewSet):
    """Community-shared resume templates."""
    permission_classes = [IsAuthenticated]
    serializer_class = None  # set dynamically

    def get_serializer_class(self):
        from .serializers import CommunityTemplateSerializer
        return CommunityTemplateSerializer

    def get_queryset(self):
        from .models import CommunityTemplate
        return CommunityTemplate.objects.filter(is_approved=True)

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def retrieve(self, request, *args, **kwargs):
        """Increment use_count when template is retrieved/used."""
        instance = self.get_object()
        instance.use_count += 1
        instance.save(update_fields=['use_count'])
        serializer = self.get_serializer(instance)
        return Response(serializer.data)


class ResumeExtractProfileView(APIView):
    """
    Uses Groq AI to extract structured profile data from the user's latest
    uploaded resume file text. Returns a JSON object matching BuilderFormData
    so the frontend can auto-fill the builder form.
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            resume = request.user.resume_data
        except ResumeData.DoesNotExist:
            return Response({'error': 'No resume uploaded yet.'}, status=400)

        # Try to re-extract text from the stored file
        try:
            resume_file = resume.resume_file
            resume_file.seek(0)
            text = extract_text_from_file(resume_file)
        except Exception:
            text = ''

        # Build profile from extracted skills as a fallback
        skills_list = resume.extracted_skills or []

        if not text and not skills_list:
            return Response({'error': 'Resume text could not be extracted.'}, status=400)

        # Use Groq to extract structured data
        api_key = os.environ.get('GROQ_API_KEY')
        if api_key and text:
            client = Groq(api_key=api_key)
            prompt = f"""Extract the following information from this resume text as a valid JSON object.
Return ONLY valid JSON. Use these exact keys:
{{
  "personal": {{ "name": "", "email": "", "phone": "", "location": "", "linkedin": "", "github": "", "portfolio": "" }},
  "summary": "",
  "skills": [],
  "experience": [{{"company": "", "role": "", "duration": "", "bullets": []}}],
  "education": [{{"institution": "", "degree": "", "year": "", "gpa": ""}}],
  "projects": [{{"name": "", "description": "", "tech": ""}}],
  "certifications": [{{"name": "", "issuer": "", "year": ""}}],
  "achievements": []
}}

Resume text:
{text[:4000]}
"""
            try:
                completion = client.chat.completions.create(
                    model='llama-3.1-8b-instant',
                    messages=[{'role': 'user', 'content': prompt}],
                    temperature=0.1,
                )
                raw = completion.choices[0].message.content.strip()
                # Strip code block markers if present
                if raw.startswith('```'):
                    raw = raw.split('```')[1]
                    if raw.startswith('json'):
                        raw = raw[4:]
                parsed = json.loads(raw)
                # Save to builder profile
                profile, _ = ResumeBuilderProfile.objects.get_or_create(user=request.user)
                profile.personal_details = parsed.get('personal', {})
                profile.summary = parsed.get('summary', '')
                profile.skills = parsed.get('skills', skills_list)
                profile.experience = parsed.get('experience', [])
                profile.education = parsed.get('education', [])
                profile.projects = parsed.get('projects', [])
                profile.certifications = parsed.get('certifications', [])
                profile.achievements = parsed.get('achievements', [])
                profile.save()
                return Response(parsed)
            except Exception as e:
                print(f'AI extraction error: {e}')

        # Fallback — return just the skills
        fallback = {
            'personal': {'name': '', 'email': '', 'phone': '', 'location': '', 'linkedin': '', 'github': '', 'portfolio': ''},
            'summary': '',
            'skills': skills_list,
            'experience': [{'company': '', 'role': '', 'duration': '', 'bullets': ['']}],
            'education': [{'institution': '', 'degree': '', 'year': ''}],
            'projects': [{'name': '', 'description': '', 'tech': ''}],
            'certifications': [{'name': '', 'issuer': '', 'year': ''}],
            'achievements': [''],
        }
        return Response(fallback)

class CustomTemplateViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = CustomTemplateSerializer

    def get_queryset(self):
        return CustomTemplate.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        file = self.request.data.get('file')
        file_type = file.name.split('.')[-1].lower() if file else 'html'
        serializer.save(user=self.request.user, file_type=file_type)

    @action(detail=True, methods=['post'])
    def detect_placeholders(self, request, pk=None):
        template = self.get_object()
        fields = []
        try:
            if template.file_type == 'docx':
                doc = Document(template.file.path)
                for para in doc.paragraphs:
                    fields.extend(re.findall(r'\{\{(.*?)\}\}', para.text))
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            fields.extend(re.findall(r'\{\{(.*?)\}\}', cell.text))
            elif template.file_type == 'pdf':
                reader = PdfReader(template.file.path)
                fields = list(reader.get_fields().keys()) if reader.get_fields() else []
            elif template.file_type in ['png', 'jpg', 'jpeg']:
                # Images don't have built-in placeholders, we suggest standard fields
                fields = ['Name', 'Email', 'Phone', 'Location', 'JobTitle', 'Summary', 'Skills']
            elif template.file_type in ['html', 'md', 'json', 'txt']:
                with open(template.file.path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    fields.extend(re.findall(r'\{\{(.*?)\}\}', content))
        except Exception as e:
            return Response({'error': str(e)}, status=500)
        
        deduped = list(set([f.strip() for f in fields]))
        return Response({'placeholders': deduped})

    @action(detail=True, methods=['post'])
    def fill_template(self, request, pk=None):
        template = self.get_object()
        data = request.data.get('data', {})
        mapping = request.data.get('mapping', template.mapped_fields)
        
        try:
            if template.file_type == 'docx':
                doc = Document(template.file.path)
                def replace_in_text(text):
                    for placeholder, fieldPath in mapping.items():
                        parts = fieldPath.split('.')
                        val = data
                        for p in parts:
                            if isinstance(val, dict): val = val.get(p, '')
                            else: val = ''
                        if isinstance(val, list):
                            val = "\n".join([str(v) if not isinstance(v, dict) else json.dumps(v) for v in val])
                        text = text.replace(f"{{{{{placeholder}}}}}", str(val))
                    return text

                for para in doc.paragraphs: para.text = replace_in_text(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs: para.text = replace_in_text(para.text)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                ext = 'docx'

            elif template.file_type == 'pdf':
                reader = PdfReader(template.file.path)
                writer = PdfWriter()
                
                # Copy pages
                for page in reader.pages:
                    writer.add_page(page)
                
                payload = {}
                for placeholder, fieldPath in mapping.items():
                    parts = fieldPath.split('.')
                    val = data
                    for p in parts:
                        if isinstance(val, dict): val = val.get(p, '')
                        else: val = ''
                    payload[placeholder] = str(val)
                
                writer.update_page_form_field_values(writer.pages[0], payload)
                
                buffer = io.BytesIO()
                writer.write(buffer)
                buffer.seek(0)
                ext = 'pdf'

            elif template.file_type in ['png', 'jpg', 'jpeg']:
                img = Image.open(template.file.path)
                draw = ImageDraw.Draw(img)
                # Use a default font if possible
                try: font = ImageFont.truetype("arial.ttf", 24)
                except: font = ImageFont.load_default()

                y_offset = 50
                for placeholder, fieldPath in mapping.items():
                    parts = fieldPath.split('.')
                    val = data
                    for p in parts:
                        if isinstance(val, dict): val = val.get(p, '')
                        else: val = ''
                    
                    text = f"{placeholder}: {val}"
                    draw.text((50, y_offset), text, fill=(0, 0, 0), font=font)
                    y_offset += 40

                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                buffer.seek(0)
                ext = 'png'

            elif template.file_type in ['html', 'md', 'txt']:
                with open(template.file.path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                for placeholder, fieldPath in mapping.items():
                    parts = fieldPath.split('.')
                    val = data
                    for p in parts:
                        if isinstance(val, dict): val = val.get(p, '')
                        else: val = ''
                    content = content.replace(f"{{{{{placeholder}}}}}", str(val))
                
                ext = template.file_type
                buffer = io.BytesIO(content.encode('utf-8'))

            else:
                return Response({'error': 'Unsupported file type'}, status=400)

            # Create entry
            last_v = GeneratedResume.objects.filter(user=self.request.user).order_by('-version_number').first()
            vn = (last_v.version_number + 1) if last_v else 1
            
            gen = GeneratedResume.objects.create(
                user=self.request.user,
                custom_template=template,
                file_name=f"{template.name}_v{vn}.{ext}",
                version_number=vn,
                data_snapshot=data
            )
            gen.file.save(f"resume_v{vn}.{ext}", ContentFile(buffer.read()))
            
            return Response({
                'success': True, 
                'resume_id': gen.id, 
                'file_url': gen.file.url,
                'version': vn
            })

        except Exception as e:
            return Response({'error': str(e)}, status=500)

class GeneratedResumeViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = GeneratedResumeSerializer

    def get_queryset(self):
        return GeneratedResume.objects.filter(user=self.request.user)